#!/usr/bin/env python3
"""
Testing framework for RAG pipeline improvements.
Run this script to test query expansion, retrieval quality, and answer generation with your actual DDQ document.
"""

import os
import sys
import asyncio
import json
import logging
from typing import List, Dict, Any
from pathlib import Path
from dotenv import load_dotenv

# Load environment variables from .env file
load_dotenv()

# Add backend to path
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from embed_utils import EmbeddingManager
from rag_utils import RAGManager

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class RAGTester:
    def __init__(self):
        self.embedding_manager = None
        self.rag_manager = None
        self.test_data_dir = Path("test_data")
        
    async def setup(self):
        """Initialize managers and load test data"""
        try:
            logger.info("Initializing RAG testing framework...")
            
            # Initialize managers
            self.embedding_manager = EmbeddingManager()
            self.rag_manager = RAGManager(self.embedding_manager)
            
            # Create test data directory
            self.test_data_dir.mkdir(exist_ok=True)
            
            logger.info("RAG testing framework initialized successfully")
            
        except Exception as e:
            logger.error(f"Error initializing testing framework: {str(e)}")
            raise
    
    def load_ddq_document(self):
        """Load your actual DDQ document for testing"""
        from docx import Document
        
        # Path to your local DDQ document
        ddq_doc_path = "/Users/imaadbukhari/Downloads/DDQ.docx"
        
        try:
            # Read Word document content
            doc = Document(ddq_doc_path)
            content = ""
            
            # Extract text from all paragraphs
            for paragraph in doc.paragraphs:
                content += paragraph.text + "\n"
            
            # Also extract text from tables if any
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        content += cell.text + " "
                    content += "\n"
            
            logger.info(f"Extracted {len(content)} characters from DDQ document")
            
        except Exception as e:
            logger.error(f"Error reading DDQ document: {str(e)}")
            raise Exception(f"Could not load DDQ document: {str(e)}")
        
        # Save it in the test data directory
        doc_path = self.test_data_dir / "ddq_document.txt"
        with open(doc_path, 'w', encoding='utf-8') as f:
            f.write(content)
        
        logger.info(f"Loaded DDQ document: {doc_path}")
        return str(doc_path)
    
    async def test_query_expansion(self):
        """Test query expansion functionality"""
        logger.info("Testing query expansion...")
        
        test_queries = [
            "What is your fund size?",
            "How much do you typically invest?",
            "What sectors do you focus on?",
            "Tell me about your portfolio"
        ]
        
        for query in test_queries:
            logger.info(f"\nOriginal query: {query}")
            try:
                expanded = await self.rag_manager._expand_query(query)
                logger.info(f"Expanded queries: {expanded}")
            except Exception as e:
                logger.error(f"Error expanding query '{query}': {str(e)}")
    
    async def test_retrieval_quality(self, test_doc_path: str):
        """Test retrieval quality with DDQ document"""
        logger.info("Testing retrieval quality...")
        
        # Process and embed the test document
        test_files = [{
            'id': 'ddq_doc_1',
            'name': 'DDQ_Document.docx',
            'content': open(test_doc_path, 'r').read(),
            'mime_type': 'text/plain',
            'modified_time': '2024-01-01T00:00:00Z'
        }]
        
        await self.embedding_manager.process_and_embed_files(test_files)
        logger.info("DDQ document processed and embedded")
        
        # Test queries with expected context levels
        test_queries = [
            {
                "query": "What is your fund size?",
                "expected_level": "fund-level",
                "should_contain": ["fund size", "total capital", "fund"],
                "should_not_contain": ["portfolio company", "round size", "series"]
            },
            {
                "query": "What is your investment thesis?",
                "expected_level": "fund-level",
                "should_contain": ["thesis", "strategy", "focus"]
            },
            {
                "query": "How much do you typically invest?",
                "expected_level": "fund-level", 
                "should_contain": ["check size", "investment amount", "fund"]
            },
            {
                "query": "What sectors do you focus on?",
                "expected_level": "fund-level",
                "should_contain": ["sectors", "focus", "thesis"]
            }
        ]
        
        for test_case in test_queries:
            query = test_case["query"]
            logger.info(f"\n{'='*50}")
            logger.info(f"Testing query: {query}")
            logger.info(f"Expected level: {test_case['expected_level']}")
            logger.info(f"{'='*50}")
            
            try:
                # Test with conversation history
                conversation = [
                    {"role": "user", "content": "Hi, I'm interested in learning about your fund."},
                    {"role": "assistant", "content": "Hello! I'd be happy to help you learn about our fund. What would you like to know?"}
                ]
                
                response, sources = await self.rag_manager.answer_question(query, conversation)
                
                # Evaluate response quality
                quality_score = self._evaluate_response_quality(test_case, response)
                
                logger.info(f"Response: {response}")
                logger.info(f"Sources: {sources}")
                logger.info(f"Quality Score: {quality_score}/10")
                
                # Check for context confusion
                if test_case.get("should_not_contain"):
                    for forbidden_term in test_case["should_not_contain"]:
                        if forbidden_term.lower() in response.lower():
                            logger.warning(f"⚠️  CONTEXT CONFUSION: Response contains '{forbidden_term}' when it shouldn't!")
                
            except Exception as e:
                logger.error(f"Error testing query '{query}': {str(e)}")
    
    def _evaluate_response_quality(self, test_case: Dict[str, Any], response: str) -> int:
        """Evaluate the quality of a response based on test criteria"""
        score = 0
        response_lower = response.lower()
        
        # Check if response contains expected terms
        for term in test_case.get("should_contain", []):
            if term.lower() in response_lower:
                score += 2
        
        # Check if response avoids forbidden terms
        forbidden_terms = test_case.get("should_not_contain", [])
        if forbidden_terms:
            forbidden_score = 0
            for term in forbidden_terms:
                if term.lower() not in response_lower:
                    forbidden_score += 2
            score += forbidden_score // len(forbidden_terms)
        
        # Bonus points for structured response
        if any(marker in response for marker in ["- ", "* ", "1. ", "2. "]):
            score += 1
        
        # Bonus points for specific numbers/details
        import re
        if re.search(r'\$[\d,]+|\d+%|\d+\s*(million|billion|M|B)', response):
            score += 1
        
        return min(score, 10)  # Cap at 10
    
    async def run_comprehensive_test(self):
        """Run comprehensive test suite"""
        logger.info("Starting comprehensive RAG testing with DDQ document...")
        
        try:
            await self.setup()
            
            # Load DDQ document
            test_doc_path = self.load_ddq_document()
            
            # Test query expansion
            await self.test_query_expansion()
            
            # Test retrieval quality
            await self.test_retrieval_quality(test_doc_path)
            
            logger.info("Comprehensive testing completed!")
            
        except Exception as e:
            logger.error(f"Error in comprehensive testing: {str(e)}")
            raise

async def main():
    """Main testing function"""
    tester = RAGTester()
    await tester.run_comprehensive_test()

if __name__ == "__main__":
    asyncio.run(main())
